### install library
import subprocess
import os
def install_library(library_name):
    try:
        # 使用subprocess调用命令行来运行pip install
        subprocess.check_call(["pip", "install", library_name])
        print(f"{library_name} 安装成功")
    except subprocess.CalledProcessError:
        print(f"无法安装 {library_name}，请手动安装")
libraries_to_install = ["requests","lark-oapi","python-docx","Pillow"]  # 替换为你需要安装的库列表
for library in libraries_to_install:
    install_library(library)



### chdir
import os
CUR_FPATH = os.path.abspath(__file__)
CUR_FDIR = os.path.dirname(CUR_FPATH)
# chdir to the directory of this script
os.chdir(CUR_FDIR)


### utils
def os_system_sure(command):
    print(f"执行命令：{command}")
    result = os.system(command)
    if result != 0:
        print(f"命令执行失败：{command}")
        exit(1)
    print(f"命令执行成功：{command}")


### config
APP_ID=""
APP_SECRET=""
TOKEN="Q3c6dJG5Go3ov6xXofZcGp43nfb"
import lark_oapi as lark
from lark_oapi.api.drive.v1 import *
import time



# import os
# if not os.path.exists("feishu2md"):
#     os_system_sure("wget https://github.com/Wsine/feishu2md/releases/download/v2.1.1/feishu2md-v2.1.1-linux-amd64.tar.gz")
#     os_system_sure("tar -zxvf feishu2md-v2.1.1-linux-amd64.tar.gz")
#     os_system_sure("rm -f feishu2md-v2.1.1-linux-amd64.tar.gz")
#     os_system_sure("rm -f README.md")
# if not os.path.exists("feishu2md"):
#     print("feishu2md 可执行文件准备失败")
#     exit(1)
# os_system_sure(f'./feishu2md config --appId {APP_ID} --appSecret {APP_SECRET}')
# os_system_sure(f'./feishu2md dl "https://fvd360f8oos.feishu.cn/docx/{TOKEN}"')



def download_doc():

    # 创建client
    client = lark.Client.builder() \
        .app_id(APP_ID) \
        .app_secret(APP_SECRET) \
        .log_level(lark.LogLevel.DEBUG) \
        .build()

    def start():

        # 构造请求对象
        request: CreateExportTaskRequest = CreateExportTaskRequest.builder() \
            .request_body(ExportTask.builder()
                        .file_extension("docx")
                        .token(TOKEN)
                        .type("docx")
                        .build()) \
            .build()

        # 发起请求
        response: CreateExportTaskResponse = client.drive.v1.export_task.create(request)

        # 处理失败返回
        if not response.success():
            lark.logger.error(
                f"client.drive.v1.export_task.create failed, code: {response.code}, msg: {response.msg}, log_id: {response.get_log_id()}")
            return

        # 处理业务结果
        lark.logger.info(lark.JSON.marshal(response.data, indent=4))
        return response.data.ticket
    ticket=start()

    time.sleep(30)

    def get_token():
        print("ticket",ticket)
        # 构造请求对象
        request: GetExportTaskRequest = GetExportTaskRequest.builder() \
            .ticket(ticket) \
            .token(TOKEN) \
            .build()

        # 发起请求
        response: GetExportTaskResponse = client.drive.v1.export_task.get(request)

        # 处理失败返回
        if not response.success():
            lark.logger.error(
                f"client.drive.v1.export_task.get failed, code: {response.code}, msg: {response.msg}, log_id: {response.get_log_id()}")
            return

        # 处理业务结果
        lark.logger.info(lark.JSON.marshal(response.data, indent=4))

        return response.data.result.file_token
    file_token=get_token()

    def download():
        print("file_token",file_token)
        # 构造请求对象
        request: DownloadExportTaskRequest = DownloadExportTaskRequest.builder() \
            .file_token(file_token) \
            .build()

        # 发起请求
        response: DownloadExportTaskResponse = client.drive.v1.export_task.download(request)

        # 处理失败返回
        if not response.success():
            lark.logger.error(
                f"client.drive.v1.export_task.download failed, code: {response.code}, msg: {response.msg}, log_id: {response.get_log_id()}")
            return

        # 处理业务结果
        f = open(f"./download.docx", "wb")
        f.write(response.file.read())
        f.close()
    download()
download_doc()


import docx
from docx import *
import json
# from loguru import logger
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.parts.image import ImagePart
import xml.etree.ElementTree as ET
from PIL import Image
def docx_to_markdown(docx_file, markdown_file,finder):
    

    def handle_picture(part):
        '''
        处理保存的图片转化成md可显示的png图片
        :param part:
        :return:
        '''
        try:
            os.makedirs("img_ori", exist_ok=True)
            print(f"文件夹 img_ori 已存在或已创建")
        except OSError as e:
            print(f"创建文件夹时出错: {e.strerror}")
        try:
            os.makedirs("img_jpeg", exist_ok=True)
            print(f"文件夹 img_ori 已存在或已创建")
        except OSError as e:
            print(f"创建文件夹时出错: {e.strerror}")
        img_name = part.partname.split('/')[-1]
        
        with open(f'img_jpeg/{img_name}', "wb") as f:
            f.write(part.blob)
        return f"![图片](img_jpeg/{img_name})"

    def is_embed(para: Paragraph, doc: Document):
        root = ET.fromstring(para._element.xml)
        for elem in root.iter():
            if elem.tag.endswith("OLEObject"):
                embed_id = elem.attrib.get("Type")
                rid = None
                for imagedata_elem in root.findall(".//{urn:schemas-microsoft-com:vml}imagedata"):
                    rid = imagedata_elem.attrib.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
                if rid and embed_id == "Embed":
                    part = doc.part.related_parts[rid]
                    # logger.info(part)
                    return part
        return None




    def is_image(para: Paragraph, doc: Document):
        images = para._element.xpath('.//pic:pic')  # 获取所有图片
        for image in images:
            for img_id in image.xpath('.//a:blip/@r:embed'):  # 获取图片id
                part = doc.part.related_parts[img_id]  # 根据图片id获取对应的图片
                # logger.info(part)
                if isinstance(part, ImagePart):
                    return part
        return None


    def table_to_markdown(serial_num, document):
        table=document.tables[serial_num - 1]
        print("\ntable",len(table.rows),len(table.columns))
        if len(table.rows)==1 and len(table.columns)==1:
            print("table code",table.rows[0].cells[0].text)
            return f"``` {table.rows[0].cells[0].text}\n ```"

        markdown_rows = []
        for row in document.tables[serial_num - 1].rows:
            cells = [cell.text for cell in row.cells]
            new_cells = [cells_word.replace('\n', '<br>') for cells_word in cells]
            markdown_row = '| ' + ' | '.join(new_cells) + ' |\n'
            # 添加 Markdown 样式的行分隔符
            if len(markdown_rows) == 1:
                markdown_rows.append('|' + ' --- |' * len(cells) + '\n')
            markdown_rows.append(markdown_row)
        print("table res",''.join(markdown_rows))
        return ''.join(markdown_rows)

    def para_font_2_md(paragraph,finder):
        fontsize=None
        for inner in paragraph.iter_inner_content():
            if isinstance(inner,Run) and inner.font!=None:
                font=inner.font
                # print("para inner",inner,font.size,font.strike,font.subscript,font.bold,font.italic,font.underline,font.color.rgb)
                fontsize=inner.font.size
                # for run_inner in inner.iter_inner_content():
                #     print("run_inner",run_inner)
        ret=f"{paragraph.text}"
        if fontsize!=None:
            if fontsize>=330000:
                ret= f"# {paragraph.text}"
            elif fontsize>=220000:
                ret= f"## {paragraph.text}"
            elif fontsize>=200000:
                ret= f"### {paragraph.text}"
            elif fontsize>=180000:
                ret= f"#### {paragraph.text}"
            elif fontsize>=160000:
                ret= f"##### {paragraph.text}"
        index=finder.index(paragraph.text)
        if index!=None:
            ret=f"{index}. {ret}"

        if len(paragraph.hyperlinks)>0:
            ret=f"[{ret}]({paragraph.hyperlinks[0].address})"

        return ret+"\n\n"

    
    def docx_to_markdown(docx_name, markdown_name,finder):
        doc = Document(docx_name)

        paragraph_styles = [
            s for s in doc.styles
        ]
        print("styles", paragraph_styles)
        # logger.info(f"文档打开完成{doc}")

        with open(markdown_name, 'w', encoding='utf-8') as md_file:
            last_paragraph_was_table = False
            table_serial_number = 0
            for element in doc.element.body.iterchildren():
                # logger.debug(element)
                last_paragraph_was_table = False
                if isinstance(element, CT_Tbl):
                    table_serial_number += 1
                    table = table_to_markdown(table_serial_number, doc)
                    md_file.write(table + '\n\n')
                    last_paragraph_was_table = True
                    # logger.debug(f"Table: Converted.")
                elif isinstance(element, CT_P):
                    paragraph = Paragraph(element, doc)
                    # 图像判断
                    img_TF = is_image(paragraph, doc)
                    if img_TF:
                        md_file.write(handle_picture(img_TF))
                    # 嵌入体判断
                    # embed = is_embed(paragraph, doc)
                    # if embed:
                    #     md_file.write(
                    #         "[链接](img_ori" + embed.partname.split('/')[-1] + ')\n\n图像占位符：balabala\n\n' + handle_picture(embed))

                    # 文本操作
                    if last_paragraph_was_table:
                        md_file.write('\n')  # 在表格后加一行空行以模仿原始文档中的空白

                    tab_stops=paragraph.paragraph_format.tab_stops
                    print("\npara",paragraph.style,paragraph.text,len(tab_stops))
                    for ts in tab_stops:
                        print("tab_stop",ts.alignment,ts.position,ts.leader)

                    for inner in paragraph.iter_inner_content():
                        if isinstance(inner,Run) and inner.font!=None:
                            print("para inner",inner,inner.font.size)


                    md_file.write(para_font_2_md(paragraph,finder))
                    last_paragraph_was_table = False
                else:
                    print("unknown element",element)
    docx_to_markdown(docx_file, markdown_file,finder)


# # Example usage:
docx_file = 'download.docx'
markdown_file = 'target.md'
# docx_to_markdown(docx_file, markdown_file)

# 匹配 md文件中 数字. 开头的行
import docx
from simplify_docx import simplify
import json

# read in a document 
my_doc = docx.Document("download.docx")

# coerce to JSON using the standard options
my_doc_as_json = simplify(my_doc)
class IndexFinder():
    index_map={}

    def para_text(self,para):
        return para["VALUE"][0]["VALUE"]

    def para_with_index(self,para):
        if "style" in para and "numPr" in para["style"]:
            # print(self.para_text(para),para)
            self.index_map[self.para_text(para)]=para["style"]["numPr"]["numId"]

    def __init__(self,jsondict):
        for para in jsondict["VALUE"][0]["VALUE"]:
            if para["TYPE"]=="paragraph":
                self.para_with_index(para)
        print(json.dumps(self.index_map, indent=4))

    def index(self,text):
        if text in self.index_map:
            return self.index_map[text]
        return None


    # def find_index(self,text):



docx_to_markdown(docx_file, markdown_file,IndexFinder(my_doc_as_json))

os.system("rm -f download.docx")
os.system("rm -f feishu2md")
os.system("cp target.md ../README.md")
os.system("rm -f target.md")
os.system("rm -rf ../img_jpeg")
os.system("mv img_jpeg ../img_jpeg")
# remove *.md
os.system("rm -f *.md")